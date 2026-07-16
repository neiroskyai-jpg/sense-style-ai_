"""Сборка документов подачи ИТМО в Word (.docx).

Зачем: конкурс просит документы в Word. Берём `submission/*.md` и рендерим в .docx через python-docx
(заголовки, списки, таблицы, жирный/курсив, цитаты). Результат — в `submission/docx/`.

Требует python-docx (у фаундера установлен 1.2.0). Если нет: `pip install python-docx`.

ЗАПУСК:
    python scripts/build_submission_docx.py            # 01-описание + 02-презентация
    python scripts/build_submission_docx.py --all      # все .md из submission/
    python scripts/build_submission_docx.py 04-cv.md   # конкретный файл
    python scripts/build_submission_docx.py --all --force   # пересобрать, ЗАТЕРЕВ правки в Word

ВАЖНО: .docx правится вручную в Word (CV и мотивационное заполняются личными фактами), а .md —
источник. Поэтому файл, который правили в Word ПОСЛЕ сборки, по умолчанию не трогаем: пересборка
молча стёрла бы ручную работу. Такой файл пропускается с предупреждением; перезаписать — `--force`.
"""
from __future__ import annotations
import argparse
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "submission"
OUT = SRC / "docx"
DEFAULT = ["01-описание-проекта.md", "02-презентация.md"]

_TOK = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)|`.+?`|\[.+?\]\(.+?\))")


def _add_runs(paragraph, text: str) -> None:
    """Разложить инлайн-markdown (**жирный**, *курсив*, `код`, [ссылка](url)) на runs Word."""
    for part in _TOK.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"
        elif re.fullmatch(r"\[.+?\]\(.+?\)", part):
            m = re.match(r"\[(.+?)\]\((.+?)\)", part)
            label, url = m.group(1), m.group(2)
            paragraph.add_run(label + (f" ({url})" if url not in label else ""))
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def build(name: str):
    from docx import Document
    from docx.shared import Pt, RGBColor

    md = (SRC / name).read_text(encoding="utf-8").split("\n")
    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = "Times New Roman"; normal.size = Pt(12)

    i = 0
    while i < len(md):
        ln = md[i]
        # таблица
        if re.match(r"^\s*\|.*\|\s*$", ln) and i + 1 < len(md) and re.match(r"^\s*\|[\s:|-]+\|\s*$", md[i + 1]):
            head = [c.strip() for c in ln.strip().strip("|").split("|")]
            rows = []
            i += 2
            while i < len(md) and re.match(r"^\s*\|.*\|\s*$", md[i]):
                rows.append([c.strip() for c in md[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(head))
            table.style = "Table Grid"
            for j, h in enumerate(head):
                cell = table.rows[0].cells[j]
                cell.paragraphs[0].clear() if hasattr(cell.paragraphs[0], "clear") else None
                _add_runs(cell.paragraphs[0], h)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for j, val in enumerate(row[:len(head)]):
                    _add_runs(cells[j].paragraphs[0], val)
            doc.add_paragraph()
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", ln)
        if m:
            _add_runs(doc.add_heading(level=len(m.group(1))), m.group(2))
            i += 1
            continue
        if re.match(r"^\s*[-*]\s+", ln):
            _add_runs(doc.add_paragraph(style="List Bullet"), re.sub(r"^\s*[-*]\s+", "", ln))
            i += 1
            continue
        if re.match(r"^\s*\d+\.\s+", ln):
            _add_runs(doc.add_paragraph(style="List Number"), re.sub(r"^\s*\d+\.\s+", "", ln))
            i += 1
            continue
        if ln.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            r = p.add_run(re.sub(r"^>\s?", "", ln)); r.italic = True
            r.font.color.rgb = RGBColor(0x5D, 0x22, 0x30)
            i += 1
            continue
        if ln.strip() == "---":
            i += 1
            continue
        if ln.strip():
            _add_runs(doc.add_paragraph(), ln)
        i += 1

    OUT.mkdir(parents=True, exist_ok=True)
    dest = OUT / (Path(name).stem + ".docx")
    doc.save(str(dest))
    return dest


def _docx_text(path: Path) -> str:
    """Текст .docx (абзацы + таблицы) — для сверки «правили ли файл руками»."""
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        parts += [c.text for row in t.rows for c in row.cells]
    return "\n".join(parts).strip()


def _edited_by_hand(name: str, dest: Path) -> bool:
    """Отличается ли существующий .docx от того, что даст сборка из .md.

    По дате судить нельзя: .docx всегда новее своего .md (его же собирают ПОСЛЕ правки исходника).
    Поэтому собираем во временный файл и сравниваем текст — сборка детерминирована, значит любое
    расхождение = правки, внесённые в Word (личные факты в CV и мотивационном).
    """
    import tempfile, shutil
    with tempfile.TemporaryDirectory() as tmp:
        keep = Path(tmp) / "keep.docx"
        shutil.copy2(dest, keep)          # бережём оригинал: build() пишет по фиксированному пути
        try:
            build(name)                   # временно перезаписывает dest
            fresh = _docx_text(dest)
        finally:
            shutil.copy2(keep, dest)      # возвращаем как было — решение принимает вызывающий
        return fresh != _docx_text(keep)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("files", nargs="*")
    ap.add_argument("--all", action="store_true")
    ap.add_argument("--force", action="store_true",
                    help="пересобрать даже то, что правили в Word (затрёт ручные правки)")
    args = ap.parse_args()
    try:
        import docx  # noqa: F401
    except ImportError:
        print("Нужен python-docx: pip install python-docx", file=sys.stderr)
        return 1

    targets = sorted(p.name for p in SRC.glob("*.md")) if args.all else (args.files or DEFAULT)
    skipped = []
    for name in targets:
        src = SRC / name
        if not src.exists():
            print(f"  ! нет файла: submission/{name}", file=sys.stderr)
            continue
        # Правки, внесённые в Word (личные факты в CV), молча пересобрать = стереть их за день
        # до подачи. Поэтому трогаем только то, что совпадает со сборкой из .md.
        dest_guess = OUT / (Path(name).stem + ".docx")
        if not args.force and dest_guess.exists() and _edited_by_hand(name, dest_guess):
            skipped.append(name)
            print(f"  ~ пропуск (правился в Word): {dest_guess.relative_to(ROOT)}")
            continue
        dest = build(name)
        # Маркеры ASCII: консоль Windows в cp1251 роняет «✓»/«→» UnicodeEncodeError — сборка
        # умирала на первом же файле, хотя это ПЕРВЫЙ шаг инструкции подачи.
        print(f"OK: {name} -> {dest.relative_to(ROOT)}")
    print("\nГотовые .docx — в submission/docx/. Открой в Word, проверь и загрузи в кабинет ИТМО.")
    if skipped:
        print(f"Не тронуты (правились в Word): {', '.join(skipped)}. Пересобрать поверх — --force.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
